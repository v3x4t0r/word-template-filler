
from openpyxl import load_workbook
from docx import Document

def fill_word_template(excel_file="source.xlsx", template_file="template.docx", 
                      output_file="filled.docx"):
    """Fill Word template with Excel data, removing unused placeholders."""
    
    # Read Excel data
    wb = load_workbook(excel_file, data_only=True)
    ws = wb["Sheet1"]
    
    replacements = {"${kunde}": ws["B1"].value}
    
    # Add products (starting from row 2)
    for row in range(2, ws.max_row + 1):
        name, value = ws[f"A{row}"].value, ws[f"B{row}"].value
        if not name:
            break
        if value:
            replacements[f"${{{name}_name}}"] = name
            replacements[f"${{{name}_sum}}"] = value
    
    # Process Word template
    doc = Document(template_file)
    
    for para in list(doc.paragraphs):  # Use list() to avoid modification issues
        text = ''.join(run.text for run in para.runs)
        
        # Replace placeholders
        for placeholder, value in replacements.items():
            text = text.replace(placeholder, str(value))
        
        if "${" in text:  # Remove paragraphs with unresolved placeholders
            para._element.getparent().remove(para._element)
        else:  # Update paragraph text
            para.clear()
            para.add_run(text)
    
    doc.save(output_file)
    print(f"✔ Document created: {output_file}")

if __name__ == "__main__":
    fill_word_template()
